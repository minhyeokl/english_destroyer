import re

import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

def read_styles(input_file):
    doc = Document(input_file)
    styles = doc.styles
    style_list = []
    for style in styles:
        style_list.append(style.name)

    return style_list


   

def destroy_english(filename, safe_styles):
    doc = Document(filename)
    output_filename = ''.join(filename.split('.')[:-1]) + '-edited.docx'

    for para in doc.paragraphs:
        if re.search('[가-힇]', para.text) == None and \
           para.style.name not in safe_styles:
            if '</>' == para.text.strip():
                continue
            
            p = para._element
            p.getparent().remove(p)
            para._p = para._element = None

    doc.save(output_filename)
